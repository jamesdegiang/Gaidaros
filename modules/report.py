from docx import Document
from docx.shared import RGBColor
import re
import datetime


R = '\033[31m' # red
G = '\033[32m' # green
C = '\033[36m' # cyan
W = '\033[0m'  # white
Y = '\033[33m' # yellow


## main
def report(target):
    try:
        if target.startswith(('http://', 'https://')):
            target = target.replace('http://', '').replace('https://', '')
        else: pass
        if target.endswith('/') == True:
            target = target[:-1]
        else: pass
        file_name = target + ".txt"
        log_path = "./dumps/"
        f = open(log_path + file_name, "r")
        logs = f.readlines()

        document = Document()
        document.add_heading('Gaidaros Security Report', 0)

        header = document.add_paragraph()
        header.add_run('Report date : ').bold = True
        current_time = datetime.datetime.now() 
        today = (str(current_time).split(' '))[0]
        time = (str(current_time).split(' '))[1]
        header.add_run(today + '\n')
        header.add_run('Report time : ').bold = True
        header.add_run(time + '\n')
        header.add_run('Target : ').bold = True
        header.add_run(target)

        doc_file_path = './reports/' + target + '.docx'
        virus(logs, document)
        apache_cve(logs, document)
        site(logs, document)
        document.save(doc_file_path)
    except Exception as e:
        print('\n' + R + '[-]' + C + ' Exception : ' + W + str(e) + '\n')


## Light Scan

# 1. Apache CVE
def apache_cve(logs, document):
    apache_cve_logs = []
    for line in logs:
        if re.search('CVE-\\d{4}-\\d{4}', line):
            apache_cve_logs.append(line.strip())
        else:
            pass
    if apache_cve_logs != []:
        cve_data = ", ".join(apache_cve_logs)
        document.add_heading('Potential Apache CVE', level=1)
        apache_cve_para = document.add_paragraph()
        apache_cve_para.add_run('Risk Level : ').bold = True
        run = apache_cve_para.add_run('Medium\n')
        font = run.font
        font.color.rgb = RGBColor(0xFF, 0x80, 0x00)
        apache_cve_para.add_run('Risk Description : ').bold = True
        apache_cve_para.add_run('Below are the Web Server Apache CVE we extracted based on your detected current Apache Server version\n')
        apache_cve_para.add_run(cve_data + '\n').italic = True
        apache_cve_para.add_run('Recommendation : ').bold = True
        apache_cve_para.add_run('We recommend you and your security team to take a look at all these mentioned CVE and work out solutions as quickly as possible')
    else:
        pass
# 2. Site Vulns
def site(logs, document):
    for line in logs:
        # Redirection
        if re.search('Redirections detected on the target', line):
            document.add_heading('Redirections detected', level=1)
            site_redirection_para = document.add_paragraph()
            site_redirection_para.add_run('Risk Level : ').bold = True
            run = site_redirection_para.add_run('Medium\n')
            font = run.font
            font.color.rgb = RGBColor(0xFF, 0x80, 0x00)
            site_redirection_para.add_run('Risk Description : ').bold = True
            site_redirection_para.add_run('Without proper validation, attackers can redirect victims to phishing or malware sites, or use forwards to access unauthorized pages\n')
            for line in logs:
                if re.search('Root page / redirects to : ', line):
                    site_redirection_para.add_run(line.strip() + '\n').italic = True
                else: pass
            site_redirection_para.add_run('Recommendation : ').bold = True
            site_redirection_para.add_run('On the surface, redirecting a user to a given website or page may seem like a harmless action. However, it can be misused to have a serious impact as it gives the hacker an open door to infect your network/website and steal customer credentials. This can put your organization’s web presence, data and credibility into jeopardy. So make sure everything is under control')
        # Server
        elif re.search('Server Information is not configured to be hidden HTTP responses header :', line):
            document.add_heading('Transfer of Sensitive Information', level=1)
            site_server_para = document.add_paragraph()
            site_server_para.add_run('Risk Level : ').bold = True
            run = site_server_para.add_run('Medium\n')
            font = run.font
            font.color.rgb = RGBColor(0xFF, 0x80, 0x00)
            site_server_para.add_run('Risk Description : ').bold = True
            site_server_para.add_run('Revealing the specific software version of the server might allow the server machine to become more vulnerable to attacks against software that is known to contain security holes\n')
            site_server_para.add_run(line.strip() + '\n').italic = True
            site_server_para.add_run('Recommendation : ').bold = True
            site_server_para.add_run('Implementors SHOULD make the Server header field a configurable option.')
            # Detect Apache Outdated version
            if re.search('Outdated Apache Web Server Version detected : ', line):
                site_outdated_apache_para = document.add_paragraph()
                site_outdated_apache_para.add_run('Risk Level : ').bold = True
                run = site_outdated_apache_para.add_run('Medium\n')
                font = run.font
                font.color.rgb = RGBColor(0xFF, 0x80, 0x00)
                site_outdated_apache_para.add_run('Risk Description : ').bold = True
                site_outdated_apache_para.add_run('Any old unpatched software presents a risk, same thing happends with Apache Web Server\n')
                site_outdated_apache_para.add_run(line.strip() + '\n').italic = True
                site_outdated_apache_para.add_run('Recommendation : ').bold = True
                site_outdated_apache_para.add_run('From security perspective, It is recommended to stay up-to-date with your Apache Web Server version and pay attention to any potential threat vectors')  
        # Detect Apache Version exposure misconfiguration via 404 response
        elif  re.search('404 Reponses exposed sensitive data about Web Server', line):
            document.add_heading('Sensitive 404 reponse disclosure', level=1)
            site_404_para = document.add_paragraph()
            site_404_para.add_run('Risk Level : ').bold = True
            run = site_404_para.add_run('Medium\n')
            font = run.font
            font.color.rgb = RGBColor(0xFF, 0x80, 0x00)
            site_404_para.add_run('Risk Description : ').bold = True
            site_404_para.add_run('Any Internet user could see the disclosed information about your webserver via an unhandled 404 HTTP response\n')
            site_404_para.add_run(line.strip() + '\n').italic = True
            site_404_para.add_run('Recommendation : ').bold = True
            site_404_para.add_run('This completely leaked information about Web Server technologies and versions could lead to many unexpected threat vectors in the future. So it is always recommended to configure and hide all these types of information')
        # X-Powered-By
        elif re.search('X-Powered-By is not configured to be hidden : ', line):
            document.add_heading('Application Framework Collection disclosure', level=1)
            site_x_powered_by_para = document.add_paragraph()
            site_x_powered_by_para.add_run('Risk Level : ').bold = True
            run = site_x_powered_by_para.add_run('Medium\n')
            font = run.font
            font.color.rgb = RGBColor(0xFF, 0x80, 0x00)
            site_x_powered_by_para.add_run('Risk Description : ').bold = True
            site_x_powered_by_para.add_run('X-Powered-By HTTP header contains the information about the collection of application frameworks being run by the site\n')
            site_x_powered_by_para.add_run(line.strip() + '\n').italic = True
            site_x_powered_by_para.add_run('Recommendation : ').bold = True
            site_x_powered_by_para.add_run('Of course there is no direct security risk, but as David C notes, exposing an outdated (and possibly vulnerable) version of PHP may be an invitation for people to try and attack it and it could be very dangerous. Remove it from the "HTTP Response Headers" session')
        # X-AspNet-Version 
        elif re.search('X-AspNet-Version is exposed : ', line) or re.search('X-AspNetMvc-Version is exposed : ', line):
            document.add_heading('Application Framework Collection disclosure', level=1)
            site_aspnet_para = document.add_paragraph()
            site_aspnet_para.add_run('Risk Level : ').bold = True
            run = site_aspnet_para.add_run('Medium\n')
            font = run.font
            font.color.rgb = RGBColor(0xFF, 0x80, 0x00)
            site_aspnet_para.add_run('Risk Description : ').bold = True
            site_aspnet_para.add_run('The X-AspNet-Version HTTP Header broadcasts to the world what version of ASP.NET is being used by your web server so the content inside is a thing worth to consider\n')
            for line in logs:
                if re.search('X-AspNet-Version is exposed : ', line):
                    site_redirection_para.add_run(line.strip() + '\n').italic = True
                else: pass
            for line in logs:
                if re.search('X-AspNetMvc-Version is exposed : ', line):
                    site_redirection_para.add_run(line.strip() + '\n').italic = True
                else: pass
            site_aspnet_para.add_run('Recommendation : ').bold = True
            site_aspnet_para.add_run('In our point of view, disabling all ASP.NET and its MVC framework information disclosure is a wise choice since we have nothing to do with it')
        # Insecure communication & Unencrypted password submissions
        elif re.search('The network communication is not secure', line):
            document.add_heading('Unavailable site encryption', level=1)
            site_ssl_para = document.add_paragraph()
            site_ssl_para.add_run('Risk Level : ').bold = True
            run = site_ssl_para.add_run('Medium\n')
            font = run.font
            font.color.rgb = RGBColor(0xFF, 0x80, 0x00)
            site_ssl_para.add_run('Risk Description : ').bold = True
            site_ssl_para.add_run('HTTP data without SSL encryption can be intercepted by third parties to gather data passed between the two systems\n')
            site_ssl_para.add_run('The network communication is not secure\n').italic = True
            site_ssl_para.add_run('Passwords are submitted unencrypted over the network\n').italic = True
            site_ssl_para.add_run('Recommendation : ').bold = True
            site_ssl_para.add_run('provide SSL (Secure Sockets Layer) certificate for your site, which creates a secure encrypted connection between the web server and the web browser is a good choice both for security and SEO')
        # Strict-Transport-Security
        elif re.search('Strict-Transport-Security is not set for ensuring SSL Encryption', line) or re.search('Could not retrieve Strict-Transport-Security for ensuring SSL Encryption', line):
            document.add_heading('Unensured SSL Encryption')
            site_strict_ssl_para = document.add_paragraph()
            site_strict_ssl_para.add_run('Risk Level : ').bold = True
            run = site_strict_ssl_para.add_run('Medium\n')
            font = run.font
            font.color.rgb = RGBColor(0xFF, 0x80, 0x00)
            site_strict_ssl_para.add_run('Risk Description : ').bold = True
            site_strict_ssl_para.add_run('Strict-Transport-Security HTTP Security Header is not set, SSL encryption on this site is not guaranteed\n')
            site_strict_ssl_para.add_run(line.strip() + '\n').italic = True
            site_strict_ssl_para.add_run('Recommendation : ').bold = True
            site_strict_ssl_para.add_run('Proper Strict-Transport-Security protects users against passive eavesdropper and active man-in-the-middle (MITM) attacks. ... Hence, it is advisable to protect as many domains/subdomains as possible using an appropriate HSTS policy')
        # Robots.txt
        elif re.search('Found /robots.txt file', line):
            document.add_heading('Robots.txt found on server', level=1)
            site_robots_para = document.add_paragraph()
            site_robots_para.add_run('Risk Level : ').bold = True
            run = site_robots_para.add_run('Low\n')
            font = run.font
            font.color.rgb = RGBColor(0x00, 0x80, 0xFF)
            site_robots_para.add_run('Risk Description : ').bold = True
            site_robots_para.add_run('According to the robots exclusion protocol (REP), the robots.txt file is used by website developers to provide instructions about their site to indexing web robots. However, there are certain risks posed by robots.txt with the very purpose of visiting the disallowed site sections\n')
            site_robots_para.add_run(line.strip() + '\n').italic = True
            site_robots_para.add_run('Recommendation : ').bold = True
            site_robots_para.add_run('Strike a balance in your security program by weighing up your assets and give them the appropriate attention, build security in to your policies and procedures and soon those low-risk findings will be a thing of the past')
        # Sitemap.xml
        elif re.search('Found /sitemap.xml file', line):
            document.add_heading('Sitemap.xml found on server', level=1)
            site_sitemap_para = document.add_paragraph()
            site_sitemap_para.add_run('Risk Level : ').bold = True
            run = site_sitemap_para.add_run('Low\n')
            font = run.font
            font.color.rgb = RGBColor(0x00, 0x80, 0xFF)
            site_sitemap_para.add_run('Risk Description : ').bold = True
            site_sitemap_para.add_run('If you are a site that suffers particularly from scrapers, for whatever reason, then you may wish to exclude sitemap entries from your robots.txt file such that bad actors cannot find them and use them to expedite their efforts.\n')
            site_sitemap_para.add_run(line.strip() + '\n').italic = True
            site_sitemap_para.add_run('Recommendation : ').bold = True
            site_sitemap_para.add_run('It is recommended to submit sitemaps via both the GSC interface and include them inside your ./robots.txt file')
        # Insecure HTTP cookies - lack Secure flag
        elif re.search('Cookies HTTP Secure flag is not set', line):
            document.add_heading('Insecure HTTP Cookies - lack Secure flag', level=1)
            site_cookies_secure_para = document.add_paragraph()
            site_cookies_secure_para.add_run('Risk Level : ').bold = True
            run = site_cookies_secure_para.add_run('Medium\n')
            font = run.font
            font.color.rgb = RGBColor(0xFF, 0x80, 0x00)
            site_cookies_secure_para.add_run('Risk Description : ').bold = True
            site_cookies_secure_para.add_run('Since the Secure flag is not set on the cookie, the browser will send it over an unencrypted channel (plain HTTP) if such a request is made. Thus, the risk exists that an attacker will intercept the clear-text communication between the browser and the server and he will steal the cookie of the user. If this is a session cookie, the attacker could gain unauthorized access to the victim\'s web session\n')
            site_cookies_secure_para.add_run(line.strip() + '\n').italic = True
            site_cookies_secure_para.add_run('Recommendation : ').bold = True
            site_cookies_secure_para.add_run('We recommend reconfiguring the web server in order to set the flag(s) Secure to all sensitive cookies.')
        # Insecure HTTP cookies - lack HttpOnly flag
        elif re.search('Cookies HTTP HttpOnly flag is not set', line):
            document.add_heading('Insecure HTTP Cookies - lack HttpOnly flag', level=1)
            site_cookies_httponly_para = document.add_paragraph()
            site_cookies_httponly_para.add_run('Risk Level : ').bold = True
            run = site_cookies_httponly_para.add_run('Medium\n')
            font = run.font
            font.color.rgb = RGBColor(0xFF, 0x80, 0x00)
            site_cookies_httponly_para.add_run('Risk Description : ').bold = True
            site_cookies_httponly_para.add_run('Lack of the HttpOnly flag permits the browser to access the cookie from client-side scripts (ex. JavaScript, VBScript, etc). This can be exploited by an attacker in conjuction with a Cross-Site Scripting (XSS) attack in order to steal the affected cookie. If this is a session cookie, the attacker could gain unauthorized access to the victim\'s web session\n')
            site_cookies_httponly_para.add_run(line.strip() + '\n').italic = True
            site_cookies_httponly_para.add_run('Recommendation : ').bold = True
            site_cookies_httponly_para.add_run('We recommend reconfiguring the web server in order to set the flag(s) HttpOnly to all sensitive cookies.')
        # Missing both and Set-Cookie header 
        elif re.search('No Set-Cookie HTTP Header retrieved', line):
            document.add_heading('Set-Cookie HTTP Header unretrieved', level=1)
            site_cookies_para = document.add_paragraph()
            site_cookies_para.add_run('Risk Level : ').bold = True
            run = site_cookies_para.add_run('Medium\n')
            font = run.font
            font.color.rgb = RGBColor(0xFF, 0x80, 0x00)
            site_cookies_para.add_run('Risk Description : ').bold = True
            site_cookies_para.add_run('We could not retrieve your Set-Cookie HTTP Header. Lacking certain Cookie HTTP security flags could end up making your system vulnerable\n')
            site_cookies_para.add_run(line.strip() + '\n').italic = True
            site_cookies_para.add_run('Recommendation : ').bold = True
            site_cookies_para.add_run('We recommend reconfiguring the web server in order to set the flag(s) Securea, HttpOnly, and other security features to all sensitive cookies.')
        # Missing HTTP Security X-XSS-Protection Header
        elif re.search('X-XSS-Protection is not available', line):
            document.add_heading('X-XSS-Protection is not available', level=1)
            site_x_xss_protection_para = document.add_paragraph()
            site_x_xss_protection_para.add_run('Risk Level : ').bold = True
            run = site_x_xss_protection_para.add_run('Low\n')
            font = run.font
            font.color.rgb = RGBColor(0x00, 0x80, 0xFF)
            site_x_xss_protection_para.add_run('Risk Description : ').bold = True
            site_x_xss_protection_para.add_run('The X-XSS-Protection HTTP header instructs the browser to stop loading web pages when they detect reflected Cross-Site Scripting (XSS) attacks. Lack of this header exposes application users to XSS attacks in case the web application contains such vulnerability\n')
            site_x_xss_protection_para.add_run(line.strip() + '\n').italic = True
            site_x_xss_protection_para.add_run('Recommendation : ').bold = True
            site_x_xss_protection_para.add_run('We recommend setting the X-XSS-Protection header to "X-XSS-Protection: 1; mode=block".\nMore information about this issue:\nhttps://developer.mozilla.org/en-US/docs/Web/HTTP/Headers/X-XSS-Protection')
        # Missing HTTP Security X-Content-Type-Options Header 
        elif re.search('X-Content-Type-Options is not set', line):
            document.add_heading('X-Content-Type-Options is not set', level=1)
            site_x_content_type_options_para = document.add_paragraph()
            site_x_content_type_options_para.add_run('Risk Level : ').bold = True
            run = site_x_content_type_options_para.add_run('Low\n')
            font = run.font
            font.color.rgb = RGBColor(0x00, 0x80, 0xFF)
            site_x_content_type_options_para.add_run('Risk Description : ').bold = True
            site_x_content_type_options_para.add_run('The HTTP X-Content-Type-Options header is addressed to Internet Explorer browser and prevents it from reinterpreting the content of a web page (MIME-sniffing) and thus overriding the value of the Content-Type header). Lack of this header could lead to attacks such as Cross-Site Scripting or phishing\n')
            site_x_content_type_options_para.add_run(line.strip() + '\n').italic = True
            site_x_content_type_options_para.add_run('Recommendation : ').bold = True
            site_x_content_type_options_para.add_run('We recommend setting the X-Content-Type-Options header to "X-Content-Type-Options: nosniff".\nMore information about this issue:\nhttps://developer.mozilla.org/en-US/docs/Web/HTTP/Headers/X-Content-Type-Options')
        # Missing HTTP Security X-Frame-Options Header
        elif re.search('The anti-clickjacking X-Frame-Options is unavailable', line) or re.search('The anti-clickjacking X-Frame-Options is unretrievable', line):
            document.add_heading('Unavailable anti-clickjacking X-Frame-Options', level=1)
            site_x_frame_options_para = document.add_paragraph()
            site_x_frame_options_para.add_run('Risk Level : ').bold = True
            run = site_x_frame_options_para.add_run('Low\n')
            font = run.font
            font.color.rgb = RGBColor(0x00, 0x80, 0xFF)
            site_x_frame_options_para.add_run('Risk Description : ').bold = True
            site_x_frame_options_para.add_run('The X-Frame-Options HTTP response header can be used to indicate whether or not a browser should be allowed to render a page inside a frame or iframe. Sites can use this to avoid clickjacking attacks, by ensuring that their content is not embedded into other sites\n')
            site_x_frame_options_para.add_run(line.strip() + '\n').italic = True
            site_x_frame_options_para.add_run('Recommendation : ').bold = True
            site_x_frame_options_para.add_run('Configure your web server to include an X-Frame-Options header and a CSP header with frame-ancestors directive. Consult Web references for more information about the possible values for this header')
        # ETags Server leaks
        elif re.search('Etags could be leaking Server inodes', line):
            document.add_heading('Etags and Server inodes', level=1)
            site_etags_para = document.add_paragraph()
            site_etags_para.add_run('Risk Level : ').bold = True
            run = site_etags_para.add_run('Low\n')
            font = run.font
            font.color.rgb = RGBColor(0x00, 0x80, 0xFF)
            site_etags_para.add_run('Risk Description : ').bold = True
            site_etags_para.add_run('An inode is a data structure used by the Linux file system. Every file and directory has an inode which stores its name, size and other data. Every inode has a number which uniquely identifies it. There is no benefits in showing this type of information\n')
            site_etags_para.add_run(line.strip() + '\n').italic = True
            site_etags_para.add_run('Recommendation : ').bold = True
            site_etags_para.add_run('The ETag is an identifier which should uniquely identify a file on the webserver, and the inode number is a number which uniquely identifies a file on the filesystem, so it seemed to make sense to use one for the other')
        # Missing HTTP Security Content-Security-Policy Header
        elif re.search('Missing Content-Security-Policy directive', line) or re.search('Could not retrieve Content-Security-Policy', line):
            document.add_heading('Missing Content-Security-Policy directive', level=1)
            site_content_security_policy_options_para = document.add_paragraph()
            site_content_security_policy_options_para.add_run('Risk Level : ').bold = True
            run = site_content_security_policy_options_para.add_run('Low\n')
            font = run.font
            font.color.rgb = RGBColor(0x00, 0x80, 0xFF)
            site_content_security_policy_options_para.add_run('Risk Description : ').bold = True
            site_content_security_policy_options_para.add_run('Apache HTTP server in certain configurations allows remote attackers to obtain sensitive information via the ETag header, which reveals the inode number, or multipart MIME boundary, which reveals child proccess IDs (PID)\n')
            site_content_security_policy_options_para.add_run(line.strip() + '\n').italic = True
            site_content_security_policy_options_para.add_run('Recommendation : ').bold = True
            site_content_security_policy_options_para.add_run('For Apache Web Server, You should try "apache-disable-inode-etag-generation" and "apache-patch-inode-leak-openbsd"')
        # Missing HTTP Security Access-Control-Allow-Origin Header
        elif re.search('Access-Control-Allow-Origin is not set', line):
            document.add_heading('Missing CORS Security Policies', level=1)
            site_cors_para = document.add_paragraph()
            site_cors_para.add_run('Risk Level : ').bold = True
            run = site_cors_para.add_run('Low\n')
            font = run.font
            font.color.rgb = RGBColor(0x00, 0x80, 0xFF)
            site_cors_para.add_run('Risk Description : ').bold = True
            site_cors_para.add_run('CORS headers come into play when a client makes a cross-origin request. In that case, the server must indicate that it allows the cross-origin operation otherwise the browser will reject the request. The two important points are that the target server must allow the operation and the client’s browser enforces it. This is a security feature as it protects the user by not letting random websites fetch data from sites he is logged in\n')
            site_cors_para.add_run(line.strip() + '\n').italic = True
            site_cors_para.add_run('Recommendation : ').bold = True
            site_cors_para.add_run('Access-Control-Allow-Origin: * is totally safe to add to any resource, unless that resource contains private data protected by something other than standard credentials. Standard credentials are cookies, HTTP basic auth, and TLS client certificates')
        # Missing HTTP Security Referrer-Policy Header
        elif re.search('Referrer-Policy is not set, site information is not under control', line) or re.search('Could not retrieve Referrer-Policy', line):
            document.add_heading('Etags and Server inodes leakage', level=1)
            site_referrer_policy_para = document.add_paragraph()
            site_referrer_policy_para.add_run('Risk Level : ').bold = True
            run = site_referrer_policy_para.add_run('Low\n')
            font = run.font
            font.color.rgb = RGBColor(0x00, 0x80, 0xFF)
            site_referrer_policy_para.add_run('Risk Description : ').bold = True
            site_referrer_policy_para.add_run('The Referrer-Policy HTTP header controls how much referrer information (sent via the Referer header) should be included with requests\n')
            site_referrer_policy_para.add_run(line.strip() + '\n').italic = True
            site_referrer_policy_para.add_run('Recommendation : ').bold = True
            site_referrer_policy_para.add_run('Some of the best practices for securing your Web Server is to set your Referrer-Policy and use the referrer in incoming requests')
        # Missing HTTP Security Permissions-Policy Header
        elif re.search('Permissions-Policy is not set, site APIs are not under-control', line) or re.search('Could not retrieve Permissions-Policy', line):
            document.add_heading('Missing Permissions-Policy - APIs are not under-control', level=1)
            site_permissions_policy_para = document.add_paragraph()
            site_permissions_policy_para.add_run('Risk Level : ').bold = True
            run = site_permissions_policy_para.add_run('Low\n')
            font = run.font
            font.color.rgb = RGBColor(0x00, 0x80, 0xFF)
            site_permissions_policy_para.add_run('Risk Description : ').bold = True
            site_permissions_policy_para.add_run('The Permissions-Policy HTTP header replaces the existing Feature-Policy header for controlling delegation of permissions and powerful features. The header uses a structured syntax, and allows sites to more tightly restrict which origins can be granted access to features\n')
            site_permissions_policy_para.add_run(line.strip() + '\n').italic = True
            site_permissions_policy_para.add_run('Recommendation : ').bold = True
            site_permissions_policy_para.add_run('The Permissions-Policy should be used in the response (server to client) to communicate the permissions policy that should be enforced by the client')
        else:
            pass
# 3. VirusTotal
def virus(logs, document):
    for line in logs:
        if re.search('Positives : ', line):
            if re.search('Positives : 0', line):
                pass
            elif re.search('Positives : \\d+', line):
                document.add_heading('Malicious Website Detected', level=1)
                site_x_frame_options_para = document.add_paragraph()
                site_x_frame_options_para.add_run('Risk Level : ').bold = True
                run = site_x_frame_options_para.add_run('High\n')
                font = run.font
                font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                site_x_frame_options_para.add_run('Risk Description : ').bold = True
                site_x_frame_options_para.add_run('VirusTotal API detected your site has been affected by malwares. Below is the number of Anti-malware engines which confirmed your site to be malicious\n')
                site_x_frame_options_para.add_run(line.strip() + '\n').italic = True
                site_x_frame_options_para.add_run('Recommendation : ').bold = True
                site_x_frame_options_para.add_run('It is essential to perform Malware Scan on your Web Server and Web Applications as soon as possible')
            else: pass

## OWASP Scan